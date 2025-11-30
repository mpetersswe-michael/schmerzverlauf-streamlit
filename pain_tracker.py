import io
import os
import datetime as dt
from typing import Tuple

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader


# ----------------------------
# Config & constants
# ----------------------------
DATA_FILE = "data.csv"

PAIN_SITUATIONS = [
    "Vor Einnahme",
    "Nach Einnahme",
    "Stabil",
    "Instabil"
]

DEFAULT_COLUMNS = [
    "Name",
    "Datum",
    "Schmerzintensität",
    "Bemerkung",
    "Schmerzsituation"
]


# ----------------------------
# Helpers
# ----------------------------
def load_data() -> pd.DataFrame:
    """Load CSV if exists, else create empty DataFrame with default columns."""
    if os.path.exists(DATA_FILE):
        df = pd.read_csv(DATA_FILE, encoding="utf-8")
        # Ensure required columns exist
        for col in DEFAULT_COLUMNS:
            if col not in df.columns:
                df[col] = ""
        # Normalize types
        if "Datum" in df.columns:
            # Try to parse dates, fallback to string
            try:
                df["Datum"] = pd.to_datetime(df["Datum"]).dt.date
            except Exception:
                pass
        return df[DEFAULT_COLUMNS]
    else:
        return pd.DataFrame(columns=DEFAULT_COLUMNS)


def save_data(df: pd.DataFrame) -> None:
    """Persist DataFrame to CSV, UTF-8."""
    df.to_csv(DATA_FILE, index=False, encoding="utf-8")


def append_entry(df: pd.DataFrame,
                 name: str,
                 date_val: dt.date,
                 intensity: int,
                 note: str,
                 situation: str) -> pd.DataFrame:
    """Append-only write."""
    new_row = {
        "Name": name.strip(),
        "Datum": date_val,
        "Schmerzintensität": intensity,
        "Bemerkung": note.strip(),
        "Schmerzsituation": situation
    }
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    return df


def filter_by_name(df: pd.DataFrame, name_filter: str) -> pd.DataFrame:
    """Case-insensitive filter by name; empty filter returns all."""
    if not name_filter:
        return df.copy()
    return df[df["Name"].str.contains(name_filter, case=False, na=False)].copy()


def plot_pain_over_time(df_filtered: pd.DataFrame) -> io.BytesIO:
    """Create a line chart for pain intensity over time for the filtered view.
    Returns PNG bytes buffer for embedding/printing/export.
    """
    buf = io.BytesIO()
    if df_filtered.empty:
        # Produce a blank figure with message
        fig, ax = plt.subplots(figsize=(7, 3))
        ax.text(0.5, 0.5, "Keine Daten für Diagramm", ha="center", va="center")
        ax.axis("off")
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150)
        buf.seek(0)
        plt.close(fig)
        return buf

    # Ensure sorting by date for timeline plotting
    dfp = df_filtered.copy()
    try:
        dfp["Datum"] = pd.to_datetime(dfp["Datum"])
    except Exception:
        pass
    dfp = dfp.sort_values("Datum")

    fig, ax = plt.subplots(figsize=(7, 3.5))
    ax.plot(dfp["Datum"], dfp["Schmerzintensität"], marker="o", linewidth=2, color="#1f77b4")
    ax.set_title("Schmerzverlauf über Zeit")
    ax.set_xlabel("Datum")
    ax.set_ylabel("Intensität (0–10)")
    ax.set_ylim(0, 10)
    ax.grid(True, linestyle="--", alpha=0.4)
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def build_word_document(df_filtered: pd.DataFrame,
                        chart_png: io.BytesIO,
                        name_filter: str) -> io.BytesIO:
    """Create a .docx with filtered table and chart."""
    doc = Document()
    title = f"Pain Tracking Bericht"
    subtitle = f"Filter: Name enthält '{name_filter}'" if name_filter else "Filter: keiner"

    doc.add_heading(title, level=1)
    doc.add_paragraph(subtitle)
    doc.add_paragraph(f"Generiert am {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Diagramm", level=2)
    # Save chart to a temp buffer and add to doc
    img_buf = io.BytesIO(chart_png.getvalue())
    doc.add_picture(img_buf, width=Inches(6))

    doc.add_heading("Daten (gefiltert)", level=2)
    # Add table
    if df_filtered.empty:
        doc.add_paragraph("Keine Daten vorhanden.")
    else:
        table = doc.add_table(rows=1, cols=len(DEFAULT_COLUMNS))
        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(DEFAULT_COLUMNS):
            hdr_cells[i].text = col

        for _, row in df_filtered.iterrows():
            tr = table.add_row().cells
            for i, col in enumerate(DEFAULT_COLUMNS):
                tr[i].text = str(row.get(col, ""))

    # Output buffer
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def build_pdf(df_filtered: pd.DataFrame,
              chart_png: io.BytesIO,
              name_filter: str) -> io.BytesIO:
    """Create a simple PDF with chart and a compact table using ReportLab."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, height - 50, "Pain Tracking Bericht")
    c.setFont("Helvetica", 11)
    subtitle = f"Filter: Name enthält '{name_filter}'" if name_filter else "Filter: keiner"
    c.drawString(40, height - 70, subtitle)
    c.drawString(40, height - 90, f"Generiert am {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Chart
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, height - 120, "Diagramm")
    img = ImageReader(chart_png)
    chart_w = width - 80
    chart_h = 200
    c.drawImage(img, 40, height - 120 - chart_h, width=chart_w, height=chart_h, preserveAspectRatio=True)

    # Table header
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, height - 350, "Daten (gefiltert)")

    # Table rows (compact listing; for many rows, add pagination)
    y = height - 370
    c.setFont("Helvetica", 9)
    if df_filtered.empty:
        c.drawString(40, y, "Keine Daten vorhanden.")
    else:
        # Header
        headers = DEFAULT_COLUMNS
        c.setFont("Helvetica-Bold", 9)
        c.drawString(40, y, " | ".join(headers))
        c.setFont("Helvetica", 9)
        y -= 14
        for _, row in df_filtered.iterrows():
            line = " | ".join(str(row.get(col, "")) for col in headers)
            # Wrap if necessary
            for chunk in wrap_text(line, max_chars=110):
                c.drawString(40, y, chunk)
                y -= 12
                if y < 60:
                    c.showPage()
                    y = height - 60

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


def wrap_text(text: str, max_chars: int = 90):
    """Simple word-wrap for PDF lines."""
    words = text.split()
    lines = []
    line = ""
    for w in words:
        if len(line) + len(w) + 1 <= max_chars:
            line = (line + " " + w).strip()
        else:
            lines.append(line)
            line = w
    if line:
        lines.append(line)
    return lines


# ----------------------------
# Streamlit UI
# ----------------------------
st.set_page_config(page_title="Pain Tracking", layout="wide")

st.title("Pain Tracking – Hybrid-Frontend")
st.caption("Append-only, filterbar, druckbar, mit Word-Export und vier Schmerzsituationen.")

# Load or init data
df = load_data()

# --- Input form
with st.form(key="input_form", clear_on_submit=False):
    st.subheader("Neuer Eintrag")
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        name = st.text_input("Name", value="", placeholder="Patientenname")
    with col2:
        date_val = st.date_input("Datum", value=dt.date.today())
    with col3:
        intensity = st.slider("Schmerzintensität (0–10)", min_value=0, max_value=10, value=5)

    situation = st.radio(
        "Schmerzsituation:",
        PAIN_SITUATIONS,
        horizontal=True
    )

    note = st.text_area("Bemerkung (optional)", height=80)

    submit = st.form_submit_button("Speichern (append-only)")
    if submit:
        # Self-checks
        errors = []
        if not name.strip():
            errors.append("Name ist erforderlich.")
        if situation not in PAIN_SITUATIONS:
            errors.append("Ungültige Schmerzsituation.")
        if errors:
            st.error("Bitte korrigieren:\n- " + "\n- ".join(errors))
        else:
            df = append_entry(df, name, date_val, intensity, note, situation)
            save_data(df)
            st.success("Eintrag gespeichert ✅ (append-only)")

st.divider()

# --- Filters & view
st.subheader("Datenansicht und Diagramm")
filter_name = st.text_input("Filter nach Name (Teilstring, optional)", value="")
df_filtered = filter_by_name(df, filter_name)

# Two columns: table and chart
c_table, c_chart = st.columns([3, 2])

with c_table:
    st.markdown("**Gefilterte Tabelle**")
    # Scrollable dataframe; wide mode ensures horizontal/vertical scrolling
    st.dataframe(df_filtered, use_container_width=True, height=380)

    # CSV download
    csv_bytes = df_filtered.to_csv(index=False).encode("utf-8")
    st.download_button(
        "CSV herunterladen",
        data=csv_bytes,
        file_name=f"pain_tracking_{dt.date.today()}.csv",
        mime="text/csv"
    )

with c_chart:
    st.markdown("**Diagramm (Schmerzverlauf)**")
    chart_png = plot_pain_over_time(df_filtered)
    st.image(chart_png, caption="Linienchart der Schmerzintensität", use_column_width=True)

st.divider()

# --- Exports & printing
st.subheader("Export und Drucken")

col_exp1, col_exp2, col_print = st.columns([1, 1, 1])

with col_exp1:
    # Word export
    word_buf = build_word_document(df_filtered, chart_png, filter_name)
    st.download_button(
        "Word (.docx) herunterladen",
        data=word_buf,
        file_name=f"pain_report_{dt.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

with col_exp2:
    # PDF export (diagram + compact table)
    pdf_buf = build_pdf(df_filtered, chart_png, filter_name)
    st.download_button(
        "PDF herunterladen",
        data=pdf_buf,
        file_name=f"pain_report_{dt.date.today()}.pdf",
        mime="application/pdf"
    )

with col_print:
    # Browser print tips
    st.markdown("**Drucken (Browser):**")
    st.write("- Windows: Strg+P")
    st.write("- macOS: Cmd+P")
    st.write("Wähle 'Hintergrundgrafiken drucken', damit das Diagramm sichtbar bleibt.")

st.caption("Hinweis: Word/PDF fassen Diagramm und gefilterte Daten zusammen. CSV bleibt der Rohdaten-Export.")



























